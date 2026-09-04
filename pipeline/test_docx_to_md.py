"""
pytest for pipeline/docx_to_md.py — synthetic Word documents built with
python-docx in the shape of the user's weekly documents (original text).

    python -m pytest pipeline -q
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

import build_week as bw  # noqa: E402
import docx_to_md as d2m  # noqa: E402


def _p(doc, text, style=None, verse=False):
    p = doc.add_paragraph(text, style=style)
    if verse:                       # the documents indent verse lines on both sides
        p.paragraph_format.left_indent = Pt(30)
        p.paragraph_format.right_indent = Pt(30)
    return p


def _multi_text_docx(path: Path) -> Path:
    doc = Document()
    _p(doc, "WEEK 3: PRESENT SUBJUNCTIVE 1", "Heading 1")
    _p(doc, "Readings:")
    _p(doc, "Fabulae Syrae, Cap. XXVII: 1. Mīnōs (Lines 1–49)")
    _p(doc, "Fabellae Latinae (ad Cap. XXV):")
    _p(doc, "63. Dāvus amīcum videt")
    _p(doc, "1. FABVLAE SYRAE (Cap. XXVII) — 1. Mīnōs", "Heading 2")
    _p(doc, "Textus Latīnus", "Heading 3")
    _p(doc, "Mīnōs rēx Crētam regēbat. Taurum habēbat.")
    _p(doc, "Neptūnus īrātus erat.")
    _p(doc, "Literal English Translation", "Heading 3")
    _p(doc, "King Minos was ruling Crete. He had a bull.")
    _p(doc, "Neptune was angry [Ablative Absolute].")
    _p(doc, "3. FABELLAE LATINAE (Cap. XXV)", "Heading 2")
    _p(doc, "63. Dāvus amīcum videt (ad Cap. XXV)", "Heading 3")
    _p(doc, "Textus Latīnus", "Heading 4")
    _p(doc, 'Dāvus: "Salvē, amīce! Quid agis?"')
    _p(doc, "Literal English Translation", "Heading 4")
    _p(doc, 'Davus: "Hello, friend! How are you?"')
    _p(doc, "Grammatica: Present Subjunctive", "Heading 2")
    _p(doc, "The present subjunctive is formed …")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Conjugation"
    t.cell(1, 0).text = "amem"
    _p(doc, "Pēnsa (Exercises)", "Heading 3")
    _p(doc, "Fill in the endings: Iūlius imperat ut servus labōret.")
    doc.save(str(path))
    return path


def _fr_docx(path: Path) -> Path:
    doc = Document()
    _p(doc, "CAPITVLVM TRĪCĒSIMVM QUĀRTUM (XXXIV)", "Heading 1")
    _p(doc, "DĒ ARTE POĒTICĀ (Lines 1–138)", "Heading 2")
    _p(doc, "Pars I (Lines 1–58): Dē Lūdīs", "Heading 3")
    _p(doc, "Textus Latīnus", "Heading 4")
    _p(doc, 'Aemilia: "Quis fīlium meum vulnerāvit?" Fabia rīdet.')
    _p(doc, "Nōn ego nōbilium sedeō studiōsus equōrum;", verse=True)
    _p(doc, "cui tamen ipsa favēs vincat ut ille precor.", verse=True)
    _p(doc, "Tum omnēs tacent.")
    _p(doc, "Literal English Translation", "Heading 4")
    _p(doc, 'Aemilia: "Who wounded my son?" Fabia laughs.')
    _p(doc, "Not I sit eager for noble horses;", verse=True)
    _p(doc, "yet that he whom thou favorest may win, I pray.", verse=True)
    _p(doc, "Then all are silent.")
    _p(doc, "Pars II (Lines 59–138)", "Heading 3")
    _p(doc, "Textus Latīnus", "Heading 4")
    _p(doc, "[59] Iūlius librum prōfert.\tLegit.")
    _p(doc, "Literal English Translation", "Heading 4")
    _p(doc, "[59] Julius brings out the book. He reads.")
    _p(doc, "Grammatica et Metrica: Dē Arte Poēticā", "Heading 3")
    _p(doc, "An elegiac couplet …")
    doc.save(str(path))
    return path


def test_multi_text_document(tmp_path):
    src = _multi_text_docx(tmp_path / "Week 3 Readings - Latin Text and Literal English Translation.docx")
    out, summary = d2m.convert(src, tmp_path / "out")
    assert out.name == "week-03.md"
    md = out.read_text(encoding="utf-8")
    assert md.startswith("# Week 3 — Fabulae Syrae & Fabellae Latīnae: ")
    assert "## Fabulae Syrae 1: Mīnōs (Lines 1–49)\n" in md          # range from the Readings list
    assert "## Fabellae Latīnae 63: Dāvus amīcum videt\n" in md      # "(ad Cap. XXV)" dropped, no range
    assert "### Textus Latīnus" in md and "### Literal English Translation" in md
    assert "Grammatica" not in md and "amem" not in md and "Pēnsa" not in md and "Readings" not in md
    assert [(s["heading"], s["la"], s["en"]) for s in summary] == [
        ("Fabulae Syrae 1: Mīnōs (Lines 1–49)", 2, 2), ("Fabellae Latīnae 63: Dāvus amīcum videt", 1, 1)]
    data, report = bw.build_from_text(3, md, merges={"en": {}, "la": {}})
    assert report["mismatches"] == [] and report["warnings"] == []
    assert [u["id"] for u in data["units"]] == ["w03:minos:b1.1", "w03:minos:b1.2", "w03:minos:b2.1", "w03:fl-63:b3.1"]
    assert [u["line_no"] for u in data["units"]] == [1, 1, None, None]
    assert data["units"][3]["unit_type"] == "turn" and data["units"][3]["speaker"] == "Dāvus"
    assert data["units"][2]["tags"] == [{"label": "ablative absolute", "la": None, "kind": "construction"}]
    assert [p["slug"] for p in data["week"]["parts"]] == ["minos", "fl-63"]


def test_fr_document_with_verse_and_markers(tmp_path):
    src = _fr_docx(tmp_path / "Week 13 Readings - Latin Text and Literal English Translation.docx")
    out, summary = d2m.convert(src, tmp_path / "out")
    md = out.read_text(encoding="utf-8")
    assert "## Pars I: Dē Lūdīs (Lines 1–58)\n" in md                # "(Lines …)" moved to the end
    assert "## Pars II (Lines 59–138)\n" in md
    assert "Nōn ego nōbilium sedeō studiōsus equōrum;\\\ncui tamen ipsa favēs vincat ut ille precor.\\\n" in md
    assert "[59] Iūlius librum prōfert. Legit." in md               # tab → space, marker verbatim
    assert "Metrica" not in md and "elegiac" not in md
    assert summary[0]["verse"] == 1 and summary[0]["marked"] is False and summary[1]["marked"] is True
    data, report = bw.build_from_text(13, md, merges={"en": {}, "la": {}})
    assert report["mismatches"] == [] and report["warnings"] == []
    us = data["units"]
    assert [u["id"] for u in us] == ["w13:b1.1", "w13:b1.2", "w13:b2.1", "w13:b2.2", "w13:b3.1", "w13:59.1", "w13:59.2"]
    assert [u["unit_type"] for u in us] == ["sentence", "sentence", "verse", "verse", "sentence", "sentence", "sentence"]
    assert us[0]["la"] == 'Aemilia: "Quis fīlium meum vulnerāvit?"' and us[0]["speaker"] is None
    assert us[3]["la"] == "cui tamen ipsa favēs vincat ut ille precor."
    assert [u["line_no"] for u in us] == [1, 1, None, None, None, 59, 59]
    assert [(p["part"], p["lines"]) for p in data["week"]["parts"]] == [("Pars I: Dē Lūdīs", "1–58"), ("Pars II", "59–138")]


def test_week_number_and_cli(tmp_path):
    assert d2m.week_number(Path("Week 9_ Capitulum XXXI - x.docx")) == 9
    assert d2m.week_number(Path("Week 14 Readings.docx")) == 14
    with pytest.raises(ValueError):
        d2m.week_number(Path("Readings.docx"))
    src = _fr_docx(tmp_path / "Week 13 Readings.docx")
    assert d2m.main([str(src), "--out", str(tmp_path / "o")]) == 0
    assert (tmp_path / "o" / "week-13.md").exists()
    assert d2m.main([str(tmp_path / "Week 4 missing.docx"), "--out", str(tmp_path / "o")]) == 1


REAL = Path(__file__).resolve().parent.parent / "source" / "docx" / "Week 1 Capitulum XXV - Latin Text and Literal English Translation.docx"


def test_week1_docx_reproduces_the_reference_units(tmp_path):
    """The converted Week 1 document builds the same 93 units as source/week-01.md."""
    ref = REAL.parent.parent / "week-01.md"
    if not (REAL.exists() and ref.exists()):
        pytest.skip("Week 1 documents not present")
    out, _ = d2m.convert(REAL, tmp_path)
    a, _ = bw.build_from_text(1, out.read_text(encoding="utf-8"))
    b, _ = bw.build_from_text(1, ref.read_text(encoding="utf-8"))
    assert len(a["units"]) == 93
    assert [(u["id"], u["la"], u["en_raw"]) for u in a["units"]] == [(u["id"], u["la"], u["en_raw"]) for u in b["units"]]
